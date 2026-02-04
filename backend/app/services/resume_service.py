# backend/app/services/resume_service.py
"""
Enhanced Resume Service with OCR and AI parsing
"""

import os
import uuid
import shutil
import tempfile
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple, BinaryIO
from sqlalchemy.orm import Session
from sqlalchemy.exc import SQLAlchemyError
import pdf2image
from PIL import Image
import pytesseract
from docx import Document
import PyPDF2
import requests
from io import BytesIO

from app.models.models import (
    Resume, Candidate, CandidateSkill, ParsedField,
    Job, JobType, JobStatus
)
from app.core.config import settings
from app.core.logging import logger
from app.services.ai_service import AIService
from app.workers.background import process_resume_upload


class ResumeService:
    """Production-grade resume processing service"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png'}
    OCR_LANGUAGES = ['eng', 'fra', 'deu', 'spa', 'ita']  # Supported OCR languages
    
    @staticmethod
    def process_resume_upload(
        file_content: Optional[bytes] = None,
        file_path: Optional[str] = None,
        file_url: Optional[str] = None,
        organization_id: uuid.UUID = None,
        user_id: uuid.UUID = None,
        db: Session = None
    ) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
        """
        Process resume upload from various sources
        """
        try:
            # Determine source and get file
            if file_content:
                file_obj = BytesIO(file_content)
                file_name = f"resume_{uuid.uuid4()}.pdf"
                file_type = ResumeService._detect_file_type(file_content)
                
            elif file_path and os.path.exists(file_path):
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                file_obj = BytesIO(file_content)
                file_name = os.path.basename(file_path)
                file_type = ResumeService._detect_file_type(file_content)
                
            elif file_url:
                # Download from URL
                response = requests.get(file_url, timeout=30)
                if response.status_code != 200:
                    return None, f"Failed to download from URL: {response.status_code}"
                
                file_content = response.content
                file_obj = BytesIO(file_content)
                file_name = file_url.split('/')[-1] or f"resume_{uuid.uuid4()}.pdf"
                file_type = ResumeService._detect_file_type(file_content)
                
            else:
                return None, "No valid file source provided"
            
            # Validate file
            if not file_type:
                return None, "Unsupported file type"
            
            # Generate IDs
            resume_id = uuid.uuid4()
            candidate_id = uuid.uuid4()
            job_id = uuid.uuid4()
            
            # Save file to storage
            storage_path = ResumeService._save_to_storage(
                file_content, resume_id, file_name, file_type
            )
            
            if not storage_path:
                return None, "Failed to save file"
            
            # Create resume record
            resume = Resume(
                id=resume_id,
                candidate_id=candidate_id,
                file_name=file_name,
                file_url=storage_path,
                file_type=file_type,
                uploaded_at=datetime.utcnow()
            )
            
            # Create candidate placeholder
            candidate_name = ResumeService._extract_name_from_filename(file_name)
            candidate = Candidate(
                id=candidate_id,
                organization_id=organization_id,
                owner_id=user_id,
                name=candidate_name,
                email=f"candidate_{int(datetime.utcnow().timestamp())}@example.com",
                status="new",
                overall_confidence=0,
                created_at=datetime.utcnow()
            )
            
            # Create parsing job
            job = Job(
                id=job_id,
                organization_id=organization_id,
                type=JobType.PARSE_RESUME,
                status=JobStatus.QUEUED,
                candidate_id=candidate_id,
                resume_id=resume_id,
                metadata={
                    "file_name": file_name,
                    "file_type": file_type,
                    "uploaded_by": str(user_id),
                    "file_size": len(file_content)
                },
                created_at=datetime.utcnow()
            )
            
            # Save to database
            db.add(candidate)
            db.add(resume)
            db.add(job)
            db.commit()
            
            # Trigger background processing
            process_resume_upload.delay(
                resume_id=str(resume_id),
                candidate_id=str(candidate_id),
                job_id=str(job_id)
            )
            
            response_data = {
                "resumeId": str(resume_id),
                "candidateId": str(candidate_id),
                "jobId": str(job_id),
                "fileName": file_name,
                "fileType": file_type,
                "fileSize": len(file_content)
            }
            
            logger.info(f"Resume uploaded: {file_name}, Size: {len(file_content)} bytes")
            return response_data, None
            
        except Exception as e:
            if db:
                db.rollback()
            logger.error(f"Failed to process resume upload: {str(e)}")
            return None, str(e)
    
    @staticmethod
    def parse_resume_content(
        resume_id: uuid.UUID,
        candidate_id: uuid.UUID,
        job_id: uuid.UUID,
        reprocess: bool = False
    ) -> bool:
        """
        Parse resume content with OCR and AI
        """
        from app.core.database import SessionLocal
        db = SessionLocal()
        
        try:
            # Get job and resume
            job = db.query(Job).filter(Job.id == job_id).first()
            resume = db.query(Resume).filter(Resume.id == resume_id).first()
            candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
            
            if not job or not resume or not candidate:
                logger.error(f"Invalid IDs: job={job_id}, resume={resume_id}, candidate={candidate_id}")
                return False
            
            # Update job status
            job.status = JobStatus.PROCESSING
            job.started_at = datetime.utcnow()
            job.progress = 10
            db.commit()
            
            # Extract text based on file type
            raw_text = ResumeService._extract_text_from_file(resume)
            
            if not raw_text or len(raw_text.strip()) < 50:
                logger.error(f"Failed to extract sufficient text from resume {resume_id}")
                job.status = JobStatus.FAILED
                job.error = "Failed to extract text from resume"
                job.completed_at = datetime.utcnow()
                db.commit()
                return False
            
            # Update job progress
            job.progress = 30
            db.commit()
            
            # Parse text using AI
            parsed_data = AIService.parse_resume_with_llm(raw_text)
            
            # Update job progress
            job.progress = 70
            db.commit()
            
            # Update candidate with parsed data
            ResumeService._update_candidate_with_parsed_data(
                db, candidate, parsed_data, resume, reprocess
            )
            
            # Update resume
            resume.raw_text = raw_text[:10000]  # Store first 10k chars
            resume.parsed_at = datetime.utcnow()
            
            # Complete job
            job.status = JobStatus.COMPLETED
            job.progress = 100
            job.completed_at = datetime.utcnow()
            job.metadata = {
                **job.metadata,
                "text_length": len(raw_text),
                "parsed_fields": len(parsed_data),
                "confidence_avg": sum(
                    parsed_data.get('confidence_scores', {}).values()
                ) / len(parsed_data.get('confidence_scores', {})) if parsed_data.get('confidence_scores') else 0
            }
            
            db.commit()
            
            logger.info(f"Resume parsed successfully: {resume.file_name}, "
                       f"Text length: {len(raw_text)}, "
                       f"Parsed fields: {len(parsed_data)}")
            return True
            
        except Exception as e:
            # Update job with error
            if 'job' in locals():
                job.status = JobStatus.FAILED
                job.error = str(e)
                job.completed_at = datetime.utcnow()
                db.commit()
            
            logger.error(f"Failed to parse resume: {str(e)}")
            return False
        finally:
            db.close()
    
    @staticmethod
    def batch_process_resumes(
        db: Session,
        organization_id: uuid.UUID,
        file_paths: List[str],
        user_id: uuid.UUID
    ) -> Dict[str, Any]:
        """
        Process multiple resumes in batch
        """
        results = {
            "success": 0,
            "failed": 0,
            "total": len(file_paths),
            "details": []
        }
        
        for file_path in file_paths:
            try:
                result, error = ResumeService.process_resume_upload(
                    file_path=file_path,
                    organization_id=organization_id,
                    user_id=user_id,
                    db=db
                )
                
                if result:
                    results["success"] += 1
                    results["details"].append({
                        "file": os.path.basename(file_path),
                        "status": "success",
                        "resumeId": result["resumeId"],
                        "candidateId": result["candidateId"]
                    })
                else:
                    results["failed"] += 1
                    results["details"].append({
                        "file": os.path.basename(file_path),
                        "status": "failed",
                        "error": error
                    })
                    
            except Exception as e:
                results["failed"] += 1
                results["details"].append({
                    "file": os.path.basename(file_path),
                    "status": "failed",
                    "error": str(e)
                })
        
        return results
    
    @staticmethod
    def validate_resume_file(file_content: bytes, file_name: str) -> Tuple[bool, Optional[str]]:
        """
        Validate resume file before processing
        """
        # Check file size (max 10MB)
        if len(file_content) > 10 * 1024 * 1024:
            return False, "File size exceeds 10MB limit"
        
        # Check file type
        file_type = ResumeService._detect_file_type(file_content)
        if not file_type:
            return False, "Unsupported file type"
        
        # Check for malicious content
        if ResumeService._contains_malicious_content(file_content):
            return False, "File contains potentially malicious content"
        
        # Check if it's a valid document
        try:
            if file_type == 'pdf':
                # Try to read PDF
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                if len(pdf_reader.pages) == 0:
                    return False, "Invalid PDF file"
                
            elif file_type == 'docx':
                # Try to read DOCX
                doc = Document(BytesIO(file_content))
                if len(doc.paragraphs) == 0:
                    return False, "Invalid DOCX file"
                    
        except Exception:
            return False, "Corrupted or invalid file"
        
        return True, None
    
    @staticmethod
    def _detect_file_type(file_content: bytes) -> Optional[str]:
        """Detect file type from content"""
        # Check magic numbers
        if file_content.startswith(b'%PDF'):
            return 'pdf'
        elif file_content.startswith(b'PK'):  # ZIP archive (DOCX is a ZIP)
            # Check for DOCX structure
            try:
                import zipfile
                with zipfile.ZipFile(BytesIO(file_content)) as zf:
                    if '[Content_Types].xml' in zf.namelist():
                        return 'docx'
            except:
                pass
        elif file_content.startswith(b'\xFF\xD8\xFF'):  # JPEG
            return 'jpg'
        elif file_content.startswith(b'\x89PNG\r\n\x1A\n'):  # PNG
            return 'png'
        elif b'{\\rtf1' in file_content[:100]:  # RTF
            return 'rtf'
        elif b'<!DOCTYPE html' in file_content[:100].lower():
            return 'html'
        
        # Fallback: check extension if filename is known
        return None
    
    @staticmethod
    def _extract_text_from_file(resume: Resume) -> str:
        """Extract text from resume file with OCR support"""
        try:
            file_path = None
            
            # Get file content
            if resume.file_url.startswith('http'):
                # Download from URL
                response = requests.get(resume.file_url, timeout=30)
                if response.status_code != 200:
                    return f"Failed to download: {response.status_code}"
                file_content = response.content
            else:
                # Read from local storage
                file_path = os.path.join(
                    settings.UPLOAD_DIR, 
                    resume.file_url.replace("/uploads/", "", 1)
                )
                if not os.path.exists(file_path):
                    return "File not found"
                
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            
            # Extract text based on file type
            if resume.file_type == 'pdf':
                # Try text extraction first
                text = ResumeService._extract_text_from_pdf(file_content)
                if len(text.strip()) > 100:
                    return text
                
                # Fallback to OCR
                return ResumeService._extract_text_with_ocr(file_content)
                
            elif resume.file_type == 'docx':
                return ResumeService._extract_text_from_docx(file_content)
                
            elif resume.file_type in ['jpg', 'jpeg', 'png']:
                return ResumeService._extract_text_with_ocr(file_content)
                
            else:
                # Try UTF-8 decoding
                try:
                    return file_content.decode('utf-8')
                except:
                    return file_content.decode('latin-1')
                
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            return f"Error extracting text: {str(e)}"
    
    @staticmethod
    def _extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text
            
        except Exception as e:
            logger.warning(f"PDF text extraction failed: {str(e)}")
            return ""
    
    @staticmethod
    def _extract_text_with_ocr(file_content: bytes) -> str:
        """Extract text using OCR"""
        try:
            # Convert to images
            images = []
            
            if file_content.startswith(b'%PDF'):
                # Convert PDF pages to images
                images = convert_from_bytes(file_content, dpi=300)
            else:
                # Load single image
                image = Image.open(BytesIO(file_content))
                images = [image]
            
            # Perform OCR on each image
            text = ""
            for image in images:
                # Preprocess image for better OCR
                image = ResumeService._preprocess_image(image)
                
                # Perform OCR
                page_text = pytesseract.image_to_string(
                    image, 
                    lang='+'.join(ResumeService.OCR_LANGUAGES)
                )
                text += page_text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"OCR failed: {str(e)}")
            return ""
    
    @staticmethod
    def _preprocess_image(image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        if image.mode != 'L':
            image = image.convert('L')
        
        # Increase contrast
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        
        # Apply threshold
        image = image.point(lambda x: 0 if x < 128 else 255, '1')
        
        return image
    
    @staticmethod
    def _extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""
    
    @staticmethod
    def _save_to_storage(
        file_content: bytes, 
        resume_id: uuid.UUID, 
        file_name: str,
        file_type: str
    ) -> str:
        """Save file to storage"""
        try:
            # Create upload directory
            upload_dir = os.path.join(settings.UPLOAD_DIR, str(resume_id))
            os.makedirs(upload_dir, exist_ok=True)
            
            # Generate safe filename
            safe_name = ResumeService._generate_safe_filename(file_name)
            
            # Save file
            file_path = os.path.join(upload_dir, safe_name)
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            return f"/uploads/{resume_id}/{safe_name}"
            
        except Exception as e:
            logger.error(f"Failed to save file: {str(e)}")
            return None
    
    @staticmethod
    def _generate_safe_filename(filename: str) -> str:
        """Generate safe filename"""
        import re
        
        # Remove path components
        name = os.path.basename(filename)
        
        # Replace special characters
        name = re.sub(r'[^\w\-_.]', '_', name)
        
        # Truncate if too long
        if len(name) > 100:
            name, ext = os.path.splitext(name)
            name = name[:95] + ext
        
        # Add timestamp for uniqueness
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        name_without_ext, ext = os.path.splitext(name)
        name = f"{name_without_ext}_{timestamp}{ext}"
        
        return name
    
    @staticmethod
    def _contains_malicious_content(file_content: bytes) -> bool:
        """Check for potentially malicious content"""
        # Check for executable content
        executable_signatures = [
            b'MZ',  # Windows EXE
            b'#!',  # Shell script
            b'<?php',  # PHP
            b'<script',  # JavaScript in HTML
        ]
        
        for signature in executable_signatures:
            if signature in file_content[:100]:
                return True
        
        # Check for null bytes (common in binaries)
        if b'\x00' in file_content[:1000]:
            return True
        
        return False
    
    @staticmethod
    def _update_candidate_with_parsed_data(
        db: Session,
        candidate: Candidate,
        parsed_data: Dict[str, Any],
        resume: Resume,
        reprocess: bool = False
    ):
        """Update candidate with parsed data"""
        # Clear existing data if reprocessing
        if reprocess:
            db.query(CandidateSkill).filter(
                CandidateSkill.candidate_id == candidate.id
            ).delete()
            
            db.query(ParsedField).filter(
                ParsedField.candidate_id == candidate.id
            ).delete()
        
        # Update candidate fields
        candidate.name = parsed_data.get('name', candidate.name)
        candidate.email = parsed_data.get('email', candidate.email)
        candidate.phone = parsed_data.get('phone', candidate.phone)
        candidate.years_experience = parsed_data.get('years_experience', candidate.years_experience)
        candidate.current_company = parsed_data.get('current_company', candidate.current_company)
        candidate.education = parsed_data.get('education', candidate.education)
        candidate.location = parsed_data.get('location', candidate.location)
        candidate.updated_at = datetime.utcnow()
        
        # Update skills
        skills = parsed_data.get('skills', [])
        for skill in skills[:20]:  # Limit to 20 skills
            candidate_skill = CandidateSkill(
                candidate_id=candidate.id,
                skill=skill,
                confidence=parsed_data.get('confidence_scores', {}).get('skills', 0.7)
            )
            db.add(candidate_skill)
        
        # Create parsed fields
        for field_name in ['name', 'email', 'phone', 'years_experience', 
                          'current_company', 'education', 'location']:
            value = parsed_data.get(field_name)
            if value:
                parsed_field = ParsedField(
                    candidate_id=candidate.id,
                    name=field_name,
                    value=str(value),
                    confidence=parsed_data.get('confidence_scores', {}).get(field_name, 0.8) * 100,
                    raw_extraction=str(value),
                    source="ai_parsing"
                )
                db.add(parsed_field)
        
        # Add skills as parsed field
        if skills:
            parsed_field = ParsedField(
                candidate_id=candidate.id,
                name='skills',
                value=','.join(skills[:10]),
                confidence=parsed_data.get('confidence_scores', {}).get('skills', 0.7) * 100,
                raw_extraction=','.join(skills[:10]),
                source="ai_parsing"
            )
            db.add(parsed_field)
        
        # Update conversation state
        if not candidate.conversation_state or reprocess:
            candidate.conversation_state = ResumeService._create_conversation_state(parsed_data)
        
        # Calculate overall confidence
        confidence_scores = parsed_data.get('confidence_scores', {})
        if confidence_scores:
            avg_confidence = sum(confidence_scores.values()) / len(confidence_scores)
            candidate.overall_confidence = avg_confidence * 100
        else:
            candidate.overall_confidence = ResumeService._calculate_completeness_score(parsed_data)
    
    @staticmethod
    def _create_conversation_state(parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create conversation state from parsed data"""
        fields = {}
        
        field_mapping = {
            "name": "name",
            "email": "email", 
            "phone": "phone",
            "years_experience": "experience",
            "current_company": "currentCompany",
            "education": "education",
            "location": "location"
        }
        
        confidence_scores = parsed_data.get('confidence_scores', {})
        
        for parse_key, field_key in field_mapping.items():
            value = parsed_data.get(parse_key)
            confidence = confidence_scores.get(parse_key, 0.0)
            
            if value:
                fields[field_key] = {
                    "value": value,
                    "confidence": confidence,
                    "asked": False,
                    "answered": confidence > 0.5,
                    "source": "resume"
                }
            else:
                fields[field_key] = {
                    "value": None,
                    "confidence": 0.0,
                    "asked": False,
                    "answered": False,
                    "source": None
                }
        
        # Handle skills separately
        skills = parsed_data.get('skills', [])
        skills_confidence = confidence_scores.get('skills', 0.0)
        
        fields["skills"] = {
            "value": skills,
            "confidence": skills_confidence,
            "asked": False,
            "answered": bool(skills) and skills_confidence > 0.5,
            "source": "resume" if skills else None
        }
        
        return {"fields": fields}
    
    @staticmethod
    def _calculate_completeness_score(parsed_data: Dict[str, Any]) -> float:
        """Calculate completeness score for parsed data"""
        required_fields = ['name', 'email', 'skills', 'years_experience']
        optional_fields = ['phone', 'current_company', 'education', 'location']
        
        filled_required = sum(1 for field in required_fields if parsed_data.get(field))
        filled_optional = sum(1 for field in optional_fields if parsed_data.get(field))
        
        total_score = (filled_required * 1.5) + filled_optional
        max_score = (len(required_fields) * 1.5) + len(optional_fields)
        
        return (total_score / max_score * 100) if max_score > 0 else 0.0
    
    @staticmethod
    def get_resume_quality_metrics(resume: Resume) -> Dict[str, Any]:
        """Get quality metrics for a resume"""
        if not resume.raw_text:
            return {"error": "Resume not parsed yet"}
        
        text = resume.raw_text
        
        metrics = {
            "text_length": len(text),
            "word_count": len(text.split()),
            "char_count": len(text),
            "line_count": text.count('\n') + 1,
            "avg_word_length": sum(len(word) for word in text.split()) / max(len(text.split()), 1),
            "parsed_at": resume.parsed_at.isoformat() if resume.parsed_at else None,
            "file_size": os.path.getsize(
                os.path.join(settings.UPLOAD_DIR, resume.file_url.replace("/uploads/", "", 1))
            ) if not resume.file_url.startswith('http') else None
        }
        
        # Calculate readability (simple version)
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        if sentences:
            avg_sentence_length = sum(len(s.split()) for s in sentences) / len(sentences)
            metrics["avg_sentence_length"] = avg_sentence_length
            
            # Simple readability score
            if avg_sentence_length < 15:
                metrics["readability"] = "high"
            elif avg_sentence_length < 25:
                metrics["readability"] = "medium"
            else:
                metrics["readability"] = "low"
        
        return metrics