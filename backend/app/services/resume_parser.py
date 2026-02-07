# backend/app/services/resume_parser.py
"""
Robust Resume Parser with multi-strategy extraction
"""

import os
import re
import json
import uuid
from datetime import datetime, date
from typing import Dict, List, Any, Optional, Tuple, BinaryIO
import tempfile
from dataclasses import dataclass
import logging

import pdfplumber
import docx
import pytesseract
from pdf2image import convert_from_path, convert_from_bytes
from PIL import Image, ImageEnhance, ImageFilter
import PyPDF2
from dateutil.parser import parse as parse_date
import magic
from io import BytesIO
import requests
import spacy
from dateutil.relativedelta import relativedelta

# Import NLP model
try:
    nlp = spacy.load("en_core_web_sm")
except:
    spacy.cli.download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

logger = logging.getLogger(__name__)

@dataclass
class ParsedResume:
    """Structured resume data"""
    # Basic Info
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    
    # Professional
    current_title: Optional[str] = None
    current_company: Optional[str] = None
    summary: Optional[str] = None
    objective: Optional[str] = None
    
    # Work Experience
    work_experience: List[Dict[str, Any]] = None
    years_experience: Optional[int] = None
    total_experience_months: Optional[int] = None
    
    # Skills
    skills: List[str] = None
    skill_categories: Dict[str, List[str]] = None
    
    # Education
    education: List[Dict[str, Any]] = None
    degree: Optional[str] = None
    university: Optional[str] = None
    graduation_year: Optional[int] = None
    
    # Location
    location: Optional[str] = None
    city: Optional[str] = None
    country: Optional[str] = None
    
    # Links
    linkedin_url: Optional[str] = None
    github_url: Optional[str] = None
    portfolio_url: Optional[str] = None
    
    # Certifications
    certifications: List[str] = None
    
    # Languages
    languages: List[str] = None
    
    # Metadata
    raw_text: Optional[str] = None
    sections: Dict[str, str] = None
    confidence_scores: Dict[str, float] = None
    parsing_method: Optional[str] = None
    
    def __post_init__(self):
        if self.work_experience is None:
            self.work_experience = []
        if self.skills is None:
            self.skills = []
        if self.skill_categories is None:
            self.skill_categories = {}
        if self.education is None:
            self.education = []
        if self.certifications is None:
            self.certifications = []
        if self.languages is None:
            self.languages = []
        if self.sections is None:
            self.sections = {}
        if self.confidence_scores is None:
            self.confidence_scores = {}


class ResumeParser:
    """Production-grade resume parser with multiple extraction strategies"""
    
    # Common patterns
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    PHONE_PATTERNS = [
        r'(\+?\d{1,2}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',
        r'\+\d{1,3}[\s.-]?\d{1,14}(?:[\s.-]?\d{1,13})?'
    ]
    URL_PATTERN = r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+'
    
    # Skill dictionaries
    TECHNICAL_SKILLS = {
        'python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue',
        'node.js', 'django', 'flask', 'fastapi', 'spring', 'express', 'laravel',
        'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'docker',
        'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'ansible', 'jenkins',
        'git', 'github', 'gitlab', 'ci/cd', 'rest', 'graphql', 'microservices'
    }
    
    SOFT_SKILLS = {
        'leadership', 'communication', 'teamwork', 'problem-solving',
        'critical thinking', 'time management', 'adaptability', 'creativity',
        'collaboration', 'project management', 'agile', 'scrum'
    }
    
    def __init__(self, use_ocr: bool = True, use_nlp: bool = True):
        self.use_ocr = use_ocr
        self.use_nlp = use_nlp
        self.logger = logging.getLogger(__name__)
        
    def parse_resume(self, file_content: bytes, filename: str = None) -> ParsedResume:
        """Main entry point for parsing a resume"""
        try:
            # Detect file type
            file_type = self._detect_file_type(file_content, filename)
            self.logger.info(f"Parsing {file_type} resume: {filename}")
            
            # Extract text
            raw_text = self._extract_text(file_content, file_type)
            
            if not raw_text or len(raw_text.strip()) < 100:
                raise ValueError(f"Insufficient text extracted: {len(raw_text or '')} chars")
            
            # Parse text
            parsed = self._parse_text(raw_text)
            parsed.raw_text = raw_text[:5000]  # Store first 5k chars
            parsed.parsing_method = file_type
            
            # Calculate confidence scores
            parsed.confidence_scores = self._calculate_confidence_scores(parsed, raw_text)
            
            # Calculate years of experience
            if parsed.work_experience:
                parsed.years_experience, parsed.total_experience_months = self._calculate_experience(
                    parsed.work_experience
                )
            
            return parsed
            
        except Exception as e:
            self.logger.error(f"Failed to parse resume: {str(e)}")
            # Return minimal parsed object
            return ParsedResume(raw_text=str(file_content[:1000]) if file_content else "")
    
    def _detect_file_type(self, content: bytes, filename: str = None) -> str:
        try:
            mime = magic.from_buffer(content[:2048], mime=True)

            if mime == 'application/pdf':
                return 'pdf'
            elif mime in (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ):
                return 'docx'
            elif mime.startswith('image/'):
                return 'image'
            elif mime == 'text/plain':
                return 'txt'

            # 🔥 Extra safety: detect HTML (bad download)
            if b"<html" in content[:500].lower():
                logger.error("Downloaded HTML instead of file")
                return 'unknown'

            return 'unknown'
        except Exception:
            return 'unknown'
    
    def _extract_text(self, content: bytes, file_type: str) -> str:
        """Extract text based on file type"""
        if file_type == 'pdf':
            return self._extract_from_pdf(content)
        elif file_type == 'docx':
            return self._extract_from_docx(content)
        elif file_type == 'image':
            return self._extract_with_ocr(content)
        elif file_type == 'txt':
            return content.decode('utf-8', errors='ignore')
        else:
            # Try all methods
            for method in [self._extract_from_pdf, self._extract_from_docx, self._extract_with_ocr]:
                try:
                    text = method(content)
                    if text and len(text.strip()) > 100:
                        return text
                except:
                    continue
            return ""
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using multiple methods"""
        text_methods = []
        
        # Method 1: pdfplumber (best for text-based PDFs)
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if len(text.strip()) > 100:
                    return text
                text_methods.append(text)
        except Exception as e:
            self.logger.debug(f"pdfplumber failed: {e}")
        
        # Method 2: PyPDF2 (fallback)
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if len(text.strip()) > 100:
                return text
            text_methods.append(text)
        except Exception as e:
            self.logger.debug(f"PyPDF2 failed: {e}")
        
        # Method 3: OCR if needed
        if self.use_ocr and len(text_methods[0] if text_methods else '') < 100:
            ocr_text = self._extract_with_ocr(content)
            if ocr_text and len(ocr_text.strip()) > 100:
                return ocr_text
        
        # Return best result
        return max(text_methods, key=len) if text_methods else ""
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(BytesIO(content))
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " | "
                    text += "\n"
            
            return text
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    def _extract_with_ocr(self, content: bytes) -> str:
        """Extract text using OCR"""
        if not self.use_ocr:
            return ""
        
        try:
            # Check if it's a PDF or image
            if content.startswith(b'%PDF'):
                # Convert PDF to images
                images = convert_from_bytes(content, dpi=300, thread_count=4)
            else:
                # It's an image
                image = Image.open(BytesIO(content))
                images = [image]
            
            # Perform OCR on each image
            all_text = ""
            for i, img in enumerate(images):
                # Preprocess image
                processed = self._preprocess_image(img)
                
                # OCR with multiple language support
                text = pytesseract.image_to_string(
                    processed,
                    lang='eng+fra+spa+deu+ita',
                    config='--psm 3 --oem 3'
                )
                
                if text.strip():
                    all_text += text + "\n\n"
            
            return all_text
        except Exception as e:
            self.logger.error(f"OCR failed: {e}")
            return ""
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR"""
        # Convert to grayscale
        if image.mode != 'L':
            image = image.convert('L')
        
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        
        # Apply mild sharpening
        image = image.filter(ImageFilter.SHARPEN)
        
        # Remove noise
        image = image.filter(ImageFilter.MedianFilter(size=3))
        
        # Binarize
        image = image.point(lambda x: 0 if x < 140 else 255)
        
        return image
    
    def _parse_text(self, text: str) -> ParsedResume:
        """Parse extracted text into structured data"""
        parsed = ParsedResume()
        
        # Basic extraction using regex
        parsed.email = self._extract_email(text)
        parsed.phone = self._extract_phone(text)
        parsed.linkedin_url = self._extract_linkedin(text)
        parsed.github_url = self._extract_github(text)
        parsed.portfolio_url = self._extract_portfolio(text)
        
        # Extract sections
        sections = self._split_into_sections(text)
        parsed.sections = sections
        
        # Extract from each section
        summary_text = sections.get('summary', '') + sections.get('objective', '')
        experience_text = sections.get('experience', '') + sections.get('work', '') + sections.get('employment', '')
        education_text = sections.get('education', '') + sections.get('academic', '')
        skills_text = sections.get('skills', '') + sections.get('technical', '') + sections.get('competencies', '')
        
        # Extract name (from beginning of document)
        parsed.name = self._extract_name(text)
        
        # Extract location
        parsed.location = self._extract_location(text)
        if parsed.location:
            parsed.city, parsed.country = self._parse_location(parsed.location)
        
        # Extract work experience
        parsed.work_experience = self._extract_work_experience(experience_text)
        if parsed.work_experience:
            parsed.current_company, parsed.current_title = self._get_current_position(parsed.work_experience)
        
        # Extract skills
        parsed.skills = self._extract_skills(skills_text + text)
        parsed.skill_categories = self._categorize_skills(parsed.skills)
        
        # Extract education
        parsed.education = self._extract_education(education_text)
        if parsed.education:
            parsed.degree, parsed.university, parsed.graduation_year = self._get_primary_education(parsed.education)
        
        # Extract summary/objective
        parsed.summary = self._extract_summary(summary_text)
        if not parsed.summary:
            parsed.summary = self._generate_summary(parsed)
        
        # Extract certifications
        parsed.certifications = self._extract_certifications(text)
        
        # Extract languages
        parsed.languages = self._extract_languages(text)
        
        return parsed
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        emails = re.findall(self.EMAIL_PATTERN, text, re.IGNORECASE)
        if emails:
            # Return the first email that looks like a personal/professional email
            for email in emails:
                if not any(x in email.lower() for x in ['example.com', 'test.com', 'placeholder']):
                    return email.lower()
        return None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number"""
        for pattern in self.PHONE_PATTERNS:
            phones = re.findall(pattern, text)
            if phones:
                # Clean and format
                phone = re.sub(r'[^\d+]', '', phones[0])
                if 10 <= len(phone) <= 15:
                    # Format as international
                    if not phone.startswith('+'):
                        if phone.startswith('1') and len(phone) == 11:
                            phone = '+' + phone
                        elif len(phone) == 10:
                            phone = '+1' + phone
                    return phone
        return None
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name"""
        # Look for name at the beginning of document
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and len(line) < 100:  # Reasonable name length
                # Check if line looks like a name (title case, 2-4 words)
                words = line.split()
                if 2 <= len(words) <= 4:
                    # Check if most words are capitalized
                    capitalized = sum(1 for w in words if w and w[0].isupper())
                    if capitalized >= len(words) * 0.8:
                        # Exclude common headers
                        if not any(x in line.lower() for x in ['resume', 'cv', 'curriculum', 'vitae', 'phone', 'email', '@']):
                            return line
        
        # Use NLP to find person names
        if self.use_nlp:
            try:
                doc = nlp(text[:2000])  # Process first 2000 chars
                for ent in doc.ents:
                    if ent.label_ == 'PERSON':
                        name = ent.text.strip()
                        if len(name.split()) >= 2:  # At least first and last name
                            return name
            except:
                pass
        
        return None
    
    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location"""
        # Common location patterns
        location_patterns = [
            r'(\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})\b',  # City, ST
            r'(\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*(\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b',  # City, Country
            r'\b(?:based in|located in|from)\s+([^,\n]+)',
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                if isinstance(matches[0], tuple):
                    return ', '.join(matches[0]).strip()
                else:
                    return matches[0].strip()
        
        # Look for address-like patterns
        address_match = re.search(r'\d+\s+[\w\s]+,\s*[\w\s]+(?:,\s*\w+)?', text)
        if address_match:
            return address_match.group().strip()
        
        return None
    
    def _parse_location(self, location: str) -> Tuple[Optional[str], Optional[str]]:
        """Parse location into city and country"""
        parts = [p.strip() for p in location.split(',')]
        if len(parts) >= 2:
            city = parts[0]
            country = parts[-1]
            return city, country
        return None, None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL"""
        patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'https?://(?:www\.)?linkedin\.com/in/[\w-]+'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group()
                if not url.startswith('http'):
                    url = 'https://' + url
                return url
        return None
    
    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL"""
        patterns = [
            r'github\.com/[\w-]+',
            r'https?://(?:www\.)?github\.com/[\w-]+'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group()
                if not url.startswith('http'):
                    url = 'https://' + url
                return url
        return None
    
    def _extract_portfolio(self, text: str) -> Optional[str]:
        """Extract portfolio URL"""
        patterns = [
            r'\b(?:portfolio|website|personal site):?\s*(https?://[^\s]+)',
            r'\b(?:http[s]?://(?:[^\s]+\.)?(?:com|io|dev|tech|me)[^\s]*)'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1) if match.groups() else match.group()
        return None
    
    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """Split text into sections based on headers"""
        sections = {}
        
        # Common resume headers
        headers = [
            'summary', 'objective', 'experience', 'work', 'employment',
            'education', 'academic', 'skills', 'technical', 'competencies',
            'projects', 'certifications', 'awards', 'languages',
            'publications', 'references'
        ]
        
        # Create regex pattern for headers
        header_pattern = r'\n\s*(?:' + '|'.join(
            [f'{h}(?::|$)' for h in headers] + 
            [h.upper() for h in headers] +
            [h.title() for h in headers]
        ) + r')\s*\n'
        
        # Split by headers
        parts = re.split(header_pattern, text, flags=re.IGNORECASE)
        
        if len(parts) > 1:
            # First part is before first header
            sections['contact'] = parts[0].strip()
            
            # Extract headers and content
            for i in range(1, len(parts), 2):
                if i < len(parts):
                    header_match = re.search(header_pattern, text[text.find(parts[i-1]):text.find(parts[i])])
                    if header_match:
                        header = header_match.group().strip().lower().rstrip(':')
                        sections[header] = parts[i].strip()
        
        return sections
    
    def _extract_work_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries"""
        experiences = []
        
        # Split by common separators
        entries = re.split(r'\n\s*(?=\d{4}|\w+\.?\s+\d{4}|present|current)', text, flags=re.IGNORECASE)
        
        for entry in entries:
            if not entry.strip():
                continue
                
            exp = {}
            
            # Extract dates
            date_pattern = r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4})\s*[-–]\s*(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4}|present|current)'
            date_match = re.search(date_pattern, entry, re.IGNORECASE)
            
            if date_match:
                exp['start_date'] = self._parse_date(date_match.group(1))
                end_date_str = date_match.group(2).lower()
                if end_date_str in ['present', 'current']:
                    exp['end_date'] = None
                    exp['is_current'] = True
                else:
                    exp['end_date'] = self._parse_date(end_date_str)
                    exp['is_current'] = False
            
            # Extract company and title
            lines = entry.split('\n')
            if lines:
                # First non-empty line often contains title and company
                first_line = lines[0].strip()
                
                # Common patterns: "Title at Company", "Title, Company", "Company - Title"
                patterns = [
                    r'(.+?)\s+at\s+(.+)',
                    r'(.+?),\s*(.+)',
                    r'(.+?)\s*[-–]\s*(.+)'
                ]
                
                for pattern in patterns:
                    match = re.match(pattern, first_line, re.IGNORECASE)
                    if match:
                        exp['title'] = match.group(1).strip()
                        exp['company'] = match.group(2).strip()
                        break
                
                if 'title' not in exp and 'company' not in exp:
                    # Fallback: split by common separators
                    parts = re.split(r'\s+at\s+|\s*,\s*|\s*[-–]\s*', first_line)
                    if len(parts) >= 2:
                        exp['title'] = parts[0].strip()
                        exp['company'] = parts[1].strip()
            
            # Extract description (rest of the entry)
            description_lines = []
            for i, line in enumerate(lines[1:], 1):
                line = line.strip()
                if line and not re.match(date_pattern, line, re.IGNORECASE):
                    # Clean bullet points
                    line = re.sub(r'^[•\-*\u2022]\s*', '', line)
                    description_lines.append(line)
            
            if description_lines:
                exp['description'] = '\n'.join(description_lines)
            
            # Extract location if present
            location_pattern = r'\b(?:in\s+)?([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)(?:\s*,\s*([A-Z]{2}))?\b'
            location_match = re.search(location_pattern, entry)
            if location_match:
                exp['location'] = location_match.group().strip()
            
            if exp.get('company') or exp.get('title'):
                experiences.append(exp)
        
        return experiences
    
    def _parse_date(self, date_str: str) -> Optional[date]:
        """Parse date string into date object"""
        try:
            return parse_date(date_str, fuzzy=True).date()
        except:
            # Try year only
            year_match = re.search(r'\d{4}', date_str)
            if year_match:
                year = int(year_match.group())
                return date(year, 1, 1)
            return None
    
    def _get_current_position(self, experiences: List[Dict]) -> Tuple[Optional[str], Optional[str]]:
        """Get current company and title from work experience"""
        for exp in experiences:
            if exp.get('is_current', False) or (exp.get('end_date') is None):
                return exp.get('company'), exp.get('title')
        
        # If no current marked, take the most recent
        if experiences:
            latest = max(experiences, key=lambda x: x.get('start_date') or date.min)
            return latest.get('company'), latest.get('title')
        
        return None, None
    
    def _calculate_experience(self, experiences: List[Dict]) -> Tuple[Optional[int], Optional[int]]:
        """Calculate total years and months of experience"""
        if not experiences:
            return None, None
        
        total_months = 0
        
        for exp in experiences:
            start = exp.get('start_date')
            end = exp.get('end_date') or datetime.now().date()
            
            if start:
                delta = relativedelta(end, start)
                exp_months = delta.years * 12 + delta.months
                total_months += max(0, exp_months)
        
        years = total_months // 12
        return years, total_months
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        skills = set()
        
        # Add from predefined lists
        for skill in self.TECHNICAL_SKILLS.union(self.SOFT_SKILLS):
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text, re.IGNORECASE):
                skills.add(skill.lower())
        
        # Extract from skills section (often bullet points)
        bullet_items = re.findall(r'[•\-*]\s*(.+?)(?=\n[•\-*]|\n\n|$)', text, re.DOTALL)
        for item in bullet_items:
            # Clean and split
            item_clean = item.strip().lower()
            if len(item_clean) < 50:  # Reasonable skill length
                skills.add(item_clean)
        
        # Extract capitalized tech words
        tech_words = re.findall(r'\b([A-Z][a-z]+(?:\.js|\.py|\.net)?)\b', text)
        for word in tech_words:
            if word.lower() in self.TECHNICAL_SKILLS:
                skills.add(word.lower())
        
        return sorted(list(skills))
    
    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills"""
        categories = {
            'technical': [],
            'soft': [],
            'tools': [],
            'languages': [],
            'other': []
        }
        
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower in self.TECHNICAL_SKILLS:
                categories['technical'].append(skill)
            elif skill_lower in self.SOFT_SKILLS:
                categories['soft'].append(skill)
            elif any(x in skill_lower for x in ['git', 'docker', 'jenkins', 'jira', 'confluence']):
                categories['tools'].append(skill)
            elif any(x in skill_lower for x in ['english', 'spanish', 'french', 'german', 'chinese']):
                categories['languages'].append(skill)
            else:
                categories['other'].append(skill)
        
        return {k: v for k, v in categories.items() if v}
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries"""
        education = []
        
        # Split by degree patterns
        degree_patterns = [
            r'\b(?:B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|Bachelor|Master|Doctorate)\b'
        ]
        
        entries = re.split(r'\n\s*(?=' + '|'.join(degree_patterns) + r')', text, re.IGNORECASE)
        
        for entry in entries:
            if not entry.strip():
                continue
                
            edu = {}
            
            # Extract degree
            degree_match = re.search(r'\b(B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|Bachelor|Master|Doctorate)\b', 
                                   entry, re.IGNORECASE)
            if degree_match:
                edu['degree'] = degree_match.group().upper()
            
            # Extract university
            # Look for capitalized phrases that might be university names
            lines = entry.split('\n')
            if lines:
                first_line = lines[0].strip()
                # Remove degree from first line to get university
                if edu.get('degree'):
                    university = first_line.replace(edu['degree'], '', 1).strip(' ,-')
                    if university:
                        edu['university'] = university
            
            # Extract year
            year_match = re.search(r'\b(19|20)\d{2}\b', entry)
            if year_match:
                edu['graduation_year'] = int(year_match.group())
            
            # Extract field of study
            field_pattern = r'\b(?:in|of)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
            field_match = re.search(field_pattern, entry, re.IGNORECASE)
            if field_match:
                edu['field'] = field_match.group(1)
            
            if edu:
                education.append(edu)
        
        return education
    
    def _get_primary_education(self, education: List[Dict]) -> Tuple[Optional[str], Optional[str], Optional[int]]:
        """Get highest/primary education"""
        if not education:
            return None, None, None
        
        # Sort by degree level
        degree_order = {'ph.d.': 4, 'doctorate': 4, 'mba': 3, 'm.s.': 3, 'm.a.': 3, 'master': 3,
                       'b.s.': 2, 'b.a.': 2, 'bachelor': 2}
        
        def get_degree_level(deg):
            if not deg:
                return 0
            deg_lower = deg.lower()
            for key, value in degree_order.items():
                if key in deg_lower:
                    return value
            return 1
        
        highest = max(education, key=lambda x: get_degree_level(x.get('degree')))
        
        return (
            highest.get('degree'),
            highest.get('university'),
            highest.get('graduation_year')
        )
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract summary/objective"""
        if not text:
            return None
        
        # Take first 2-3 sentences
        sentences = re.split(r'[.!?]+', text)
        summary_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence.split()) >= 5:
                summary_sentences.append(sentence)
                if len(summary_sentences) >= 3:
                    break
        
        if summary_sentences:
            return ' '.join(summary_sentences) + '.'
        return None
    
    def _generate_summary(self, parsed: ParsedResume) -> Optional[str]:
        """Generate summary from parsed data"""
        parts = []
        
        if parsed.current_title:
            parts.append(f"{parsed.current_title}")
        
        if parsed.years_experience:
            parts.append(f"with {parsed.years_experience} years of experience")
        
        if parsed.skills:
            top_skills = parsed.skills[:5]
            parts.append(f"skilled in {', '.join(top_skills)}")
        
        if parts:
            return ' '.join(parts) + '.'
        return None
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certs = set()
        
        # Common certification patterns
        cert_patterns = [
            r'\b(AWS\s+Certified|Azure\s+Certified|Google\s+Cloud\s+Certified|PMP|CISSP|CEH|CCNA|CCNP|Scrum\s+Master|SAFe)\b',
            r'\b([A-Z]{2,6}P\b|\b[A-Z]{3,}\s+[A-Z]{2,})'
        ]
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    certs.add(' '.join(match).strip())
                else:
                    certs.add(match.strip())
        
        return sorted(list(certs))
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages"""
        languages = set()
        
        # Common languages
        common_langs = ['english', 'spanish', 'french', 'german', 'chinese', 
                       'hindi', 'arabic', 'portuguese', 'russian', 'japanese']
        
        for lang in common_langs:
            pattern = r'\b' + lang + r'\b'
            if re.search(pattern, text, re.IGNORECASE):
                languages.add(lang.title())
        
        return sorted(list(languages))
    
    def _calculate_confidence_scores(self, parsed: ParsedResume, raw_text: str) -> Dict[str, float]:
        """Calculate confidence scores for parsed fields"""
        scores = {}
        
        # Name confidence
        if parsed.name:
            # Check if name appears in text
            name_in_text = parsed.name.lower() in raw_text.lower()
            # Check name format
            name_words = parsed.name.split()
            valid_format = 2 <= len(name_words) <= 4
            
            scores['name'] = 0.9 if name_in_text and valid_format else 0.6
        else:
            scores['name'] = 0.0
        
        # Email confidence
        if parsed.email:
            # Validate email format
            if re.match(self.EMAIL_PATTERN, parsed.email):
                scores['email'] = 0.95
            else:
                scores['email'] = 0.5
        else:
            scores['email'] = 0.0
        
        # Phone confidence
        if parsed.phone:
            # Check if it looks like a valid phone number
            digits = sum(c.isdigit() for c in parsed.phone)
            scores['phone'] = 0.9 if digits >= 10 else 0.5
        else:
            scores['phone'] = 0.0
        
        # Work experience confidence
        if parsed.work_experience:
            # Based on number of entries and completeness
            complete_entries = sum(1 for exp in parsed.work_experience 
                                 if exp.get('company') and exp.get('title'))
            scores['work_experience'] = min(0.9, complete_entries * 0.3)
        else:
            scores['work_experience'] = 0.0
        
        # Skills confidence
        if parsed.skills:
            # Based on number of skills and presence in skill dictionaries
            known_skills = sum(1 for skill in parsed.skills 
                             if skill.lower() in self.TECHNICAL_SKILLS.union(self.SOFT_SKILLS))
            scores['skills'] = min(0.95, known_skills * 0.1 + 0.3)
        else:
            scores['skills'] = 0.0
        
        return scores